import docx

def parse_md_to_docx(filename, doc):
    with open(filename, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line == '':
                continue
            else:
                doc.add_paragraph(line)

doc = docx.Document()
doc.add_heading('Green Chain: Platform Documentation', 0)

try:
    parse_md_to_docx('docs/architecture.md', doc)
    doc.add_page_break()
    parse_md_to_docx('docs/api_reference.md', doc)
except Exception as e:
    print(f"Error parsing docs: {e}")

doc.save('GreenChain_Documentation.docx')
print("Successfully generated GreenChain_Documentation.docx")
